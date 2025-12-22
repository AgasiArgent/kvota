"""
Specification Export Service - Russian B2B Quotation System
Generates DOCX specification documents (Спецификация) from quotes
"""
import os
from decimal import Decimal
from datetime import date, datetime
from typing import Dict, List, Optional, Any
from uuid import UUID

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from num2words import num2words

from services.export_data_mapper import format_payment_terms


# ============================================================================
# NAME FORMATTING HELPERS
# ============================================================================

def format_signatory_name_short(last_name: str, first_name: str, patronymic: str = "") -> str:
    """
    Format Russian name as "Фамилия И. О." for official documents

    Args:
        last_name: Surname (фамилия) - e.g., "Ермаков"
        first_name: First name (имя) - e.g., "Иван"
        patronymic: Patronymic (отчество) - e.g., "Иванович"

    Returns:
        Formatted name string

    Examples:
        >>> format_signatory_name_short("Ермаков", "Иван", "Иванович")
        "Ермаков И. И."
        >>> format_signatory_name_short("Ермаков", "Иван", "")
        "Ермаков И."
        >>> format_signatory_name_short("", "Иван", "")
        "Иван"
    """
    if not last_name:
        # No last name - just return first name
        return first_name.strip() if first_name else ""

    # Start with last name
    result = last_name.strip()

    # Add first name initial
    if first_name and first_name.strip():
        result += f" {first_name.strip()[0]}."

    # Add patronymic initial
    if patronymic and patronymic.strip():
        result += f" {patronymic.strip()[0]}."

    return result


def format_signatory_name_full(last_name: str, first_name: str, patronymic: str = "") -> str:
    """
    Format Russian name as "Фамилия Имя Отчество" for full display

    Args:
        last_name: Surname (фамилия) - e.g., "Ермаков"
        first_name: First name (имя) - e.g., "Иван"
        patronymic: Patronymic (отчество) - e.g., "Иванович"

    Returns:
        Full formatted name string

    Examples:
        >>> format_signatory_name_full("Ермаков", "Иван", "Иванович")
        "Ермаков Иван Иванович"
        >>> format_signatory_name_full("Ермаков", "Иван", "")
        "Ермаков Иван"
    """
    parts = []
    if last_name and last_name.strip():
        parts.append(last_name.strip())
    if first_name and first_name.strip():
        parts.append(first_name.strip())
    if patronymic and patronymic.strip():
        parts.append(patronymic.strip())

    return " ".join(parts)


def _format_org_director_name(organization: Dict[str, Any]) -> str:
    """
    Format organization director name for specification export.

    Uses new separate fields if available, falls back to legacy general_director_name.

    Args:
        organization: Organization dict with director fields

    Returns:
        Formatted name string in "Фамилия И. О." format
    """
    # Try new separate fields first
    last_name = organization.get("general_director_last_name") or ""
    first_name = organization.get("general_director_first_name") or ""
    patronymic = organization.get("general_director_patronymic") or ""

    if last_name or first_name:
        # Use new fields with short format
        return format_signatory_name_short(last_name, first_name, patronymic)

    # Fall back to legacy single field
    return organization.get("general_director_name") or ""


def _format_seller_company_director_name(seller_company: Dict[str, Any]) -> str:
    """
    Format seller company director name for specification export.

    Args:
        seller_company: Seller company dict with director fields

    Returns:
        Formatted name string in "Фамилия И. О." format
    """
    last_name = seller_company.get("general_director_last_name") or ""
    first_name = seller_company.get("general_director_first_name") or ""
    patronymic = seller_company.get("general_director_patronymic") or ""

    if last_name or first_name:
        return format_signatory_name_short(last_name, first_name, patronymic)

    return ""


# ============================================================================
# CURRENCY HELPERS
# ============================================================================

def currency_name_russian(currency: str, case: str = "nominative") -> str:
    """
    Get Russian currency name in specified grammatical case

    Args:
        currency: Currency code (USD, EUR, RUB, TRY, CNY)
        case: Grammatical case (nominative, genitive, prepositional)

    Returns:
        Russian currency name in specified case
    """
    currency_names = {
        "USD": {
            "nominative": "доллар США",
            "genitive": "долларов США",
            "prepositional": "долларах США"
        },
        "EUR": {
            "nominative": "евро",
            "genitive": "евро",
            "prepositional": "евро"
        },
        "RUB": {
            "nominative": "рубль",
            "genitive": "рублей",
            "prepositional": "рублях"
        },
        "TRY": {
            "nominative": "турецкая лира",
            "genitive": "турецких лир",
            "prepositional": "турецких лирах"
        },
        "CNY": {
            "nominative": "юань",
            "genitive": "юаней",
            "prepositional": "юанях"
        }
    }

    if currency not in currency_names:
        return currency

    return currency_names[currency].get(case, currency_names[currency]["nominative"])


def number_to_russian_words(amount: Decimal, currency: str) -> str:
    """
    Convert number to Russian words (прописью)

    Args:
        amount: Decimal amount
        currency: Currency code (USD, EUR, RUB, TRY, CNY)

    Returns:
        Russian words representation
        - For whole numbers: "Сто семь тысяч долларов США"
        - For decimals: "Сто семь тысяч долларов США 33 цента"

    Example:
        >>> number_to_russian_words(Decimal("1500.50"), "USD")
        "Одна тысяча пятьсот долларов США 50 центов"
        >>> number_to_russian_words(Decimal("107000.00"), "USD")
        "Сто семь тысяч долларов США"
    """
    # Split into integer and fractional parts
    integer_part = int(amount)
    fractional_part = int(round((amount - integer_part) * 100))

    # Convert integer part to Russian words
    # num2words for Russian uses genitive case for currency
    if integer_part == 0:
        words = "Ноль"
    else:
        words = num2words(integer_part, lang='ru').capitalize()

    # Get currency name in genitive case
    currency_name = currency_name_russian(currency, "genitive")

    # Build result - add cents only if fractional part is non-zero
    if fractional_part == 0:
        result = f"{words} {currency_name}"
    else:
        # Get cents word in proper grammatical form
        cents_word = _get_cents_word(fractional_part, currency)
        cents_num_words = num2words(fractional_part, lang='ru')
        result = f"{words} {currency_name} {cents_num_words} {cents_word}"

    return result


def _get_cents_word(cents: int, currency: str) -> str:
    """
    Get the appropriate word for cents/kopecks based on number and currency

    Russian grammar rules for numbers:
    - 1: цент (singular)
    - 2-4: цента (genitive singular)
    - 5-20: центов (genitive plural)
    - 21: цент, 22-24: цента, 25-30: центов, etc.
    """
    cents_names = {
        "USD": ("цент", "цента", "центов"),
        "EUR": ("цент", "цента", "центов"),
        "RUB": ("копейка", "копейки", "копеек"),
        "TRY": ("куруш", "куруша", "курушей"),
        "CNY": ("фынь", "фыня", "фыней"),
    }

    words = cents_names.get(currency, ("цент", "цента", "центов"))

    # Russian grammar rules for number agreement
    if 11 <= cents <= 19:
        return words[2]  # plural genitive (центов)
    last_digit = cents % 10
    if last_digit == 1:
        return words[0]  # singular (цент)
    elif 2 <= last_digit <= 4:
        return words[1]  # genitive singular (цента)
    else:
        return words[2]  # plural genitive (центов)


# ============================================================================
# DATA GATHERING
# ============================================================================

async def gather_specification_data(
    quote_id: UUID,
    contract_id: UUID,
    org_id: UUID,
    supabase_client,
    warehouse_index: Optional[int] = None,
    delivery_address_id: Optional[UUID] = None,
    signatory_contact_id: Optional[UUID] = None
) -> Dict[str, Any]:
    """
    Gather all data needed for specification export

    Args:
        quote_id: Quote UUID
        contract_id: Contract UUID
        org_id: Organization UUID
        supabase_client: Supabase client instance
        warehouse_index: Legacy - Index into customer.warehouse_addresses JSONB array
        delivery_address_id: New - UUID of delivery address from customer_delivery_addresses table
        signatory_contact_id: New - UUID of signatory contact from customer_contacts table

    Returns dict with 30 template variables:
    - Contract info: contract_number, contract_date, specification_number, specification_date
    - Company info: seller_company, customer_name
    - Quote info: currency, payment_terms, delivery_days, delivery_terms
    - Products: list of products with idn, product_code, name, brand, quantity, unit_price_vat, total_price_vat
    - Totals: total_quantity, total_amount_vat, total_amount_words, vat_amount_words
    - Addresses: customer_registration_address, customer_warehouse_address
    - Signatories: seller_signatory_position, seller_signatory_name, customer_signatory_position, customer_signatory_name
    - Optional: additional_conditions
    """

    # 1. Fetch quote data
    quote_result = supabase_client.table("quotes").select("*")\
        .eq("id", str(quote_id))\
        .eq("organization_id", str(org_id))\
        .execute()

    if not quote_result.data or len(quote_result.data) == 0:
        raise ValueError(f"Quote {quote_id} not found")

    quote = quote_result.data[0]

    # 2. Fetch contract data
    contract_result = supabase_client.table("customer_contracts").select("*")\
        .eq("id", str(contract_id))\
        .eq("organization_id", str(org_id))\
        .execute()

    if not contract_result.data or len(contract_result.data) == 0:
        raise ValueError(f"Contract {contract_id} not found")

    contract = contract_result.data[0]

    # 3. Fetch customer data
    customer_result = supabase_client.table("customers").select("*")\
        .eq("id", str(quote["customer_id"]))\
        .eq("organization_id", str(org_id))\
        .execute()

    if not customer_result.data or len(customer_result.data) == 0:
        raise ValueError(f"Customer {quote['customer_id']} not found")

    customer = customer_result.data[0]

    # 4. Fetch organization data
    org_result = supabase_client.table("organizations").select("*")\
        .eq("id", str(org_id))\
        .execute()

    if not org_result.data or len(org_result.data) == 0:
        raise ValueError(f"Organization {org_id} not found")

    organization = org_result.data[0]

    # 5. Fetch quote items (products)
    items_result = supabase_client.table("quote_items").select("*")\
        .eq("quote_id", str(quote_id))\
        .order("created_at")\
        .execute()

    if not items_result.data or len(items_result.data) == 0:
        raise ValueError(f"No products found for quote {quote_id}")

    # 5a. Fetch calculation results for sales prices
    calc_results = supabase_client.table("quote_calculation_results").select("*")\
        .eq("quote_id", str(quote_id))\
        .execute()

    # Get USD to quote currency rate for fallback conversion
    usd_to_quote_rate = float(quote.get('usd_to_quote_rate', 1.0))

    # Build lookup: quote_item_id -> phase_results (prefer quote currency, fallback to USD with conversion)
    calc_results_by_item: Dict[str, Dict] = {}
    if calc_results.data:
        for calc in calc_results.data:
            item_id = calc.get("quote_item_id")
            if item_id:
                # Prefer phase_results_quote_currency (already converted to client currency)
                phase_results_quote = calc.get("phase_results_quote_currency", {}) or {}
                phase_results_usd = calc.get("phase_results", {}) or {}

                if phase_results_quote:
                    # Use quote currency version, but merge missing fields from USD
                    merged = dict(phase_results_quote)
                    required_fields = [
                        'sales_price_per_unit_with_vat', 'sales_price_total_with_vat',
                        'sales_price_per_unit_no_vat', 'sales_price_total_no_vat',
                        'vat_from_sales'
                    ]
                    for field in required_fields:
                        if field not in merged or merged.get(field) is None:
                            if field in phase_results_usd and phase_results_usd.get(field) is not None:
                                try:
                                    merged[field] = float(phase_results_usd[field]) * usd_to_quote_rate
                                except (ValueError, TypeError):
                                    pass
                    calc_results_by_item[item_id] = merged
                elif phase_results_usd:
                    # Fallback: convert USD to quote currency
                    converted = {}
                    for field in ['sales_price_per_unit_with_vat', 'sales_price_total_with_vat',
                                  'sales_price_per_unit_no_vat', 'sales_price_total_no_vat', 'vat_from_sales']:
                        if field in phase_results_usd and phase_results_usd.get(field) is not None:
                            try:
                                converted[field] = float(phase_results_usd[field]) * usd_to_quote_rate
                            except (ValueError, TypeError):
                                pass
                    calc_results_by_item[item_id] = converted
                else:
                    calc_results_by_item[item_id] = {}

    # 5b. Fetch quote calculation variables for payment terms and delivery time
    variables_result = supabase_client.table("quote_calculation_variables").select("variables")\
        .eq("quote_id", str(quote_id))\
        .execute()

    # Extract variables dict, or empty dict if not found
    calculation_variables: Dict[str, Any] = {}
    if variables_result.data and len(variables_result.data) > 0:
        calculation_variables = variables_result.data[0].get("variables", {})

    # 5c. Fetch seller company for seller signatory
    # Look up seller company by name from calculation variables
    seller_company_name = calculation_variables.get("seller_company", "")
    seller_company_data: Dict[str, Any] = {}

    if seller_company_name:
        seller_company_result = supabase_client.table("seller_companies").select("*")\
            .eq("organization_id", str(org_id))\
            .eq("name", seller_company_name)\
            .execute()

        if seller_company_result.data and len(seller_company_result.data) > 0:
            seller_company_data = seller_company_result.data[0]

    # 6. Get warehouse address
    # Priority: delivery_address_id > warehouse_index > customer.address
    customer_warehouse_address = customer.get("address", "Адрес не указан")

    if delivery_address_id:
        # New approach: fetch from customer_delivery_addresses table
        addr_result = supabase_client.table("customer_delivery_addresses").select("*")\
            .eq("id", str(delivery_address_id))\
            .eq("organization_id", str(org_id))\
            .execute()

        if addr_result.data and len(addr_result.data) > 0:
            customer_warehouse_address = addr_result.data[0].get("address", customer_warehouse_address)
    elif warehouse_index is not None:
        # Legacy approach: use index into customer.warehouse_addresses JSONB
        warehouse_addresses = customer.get("warehouse_addresses", [])
        if warehouse_addresses and warehouse_index < len(warehouse_addresses):
            warehouse_obj = warehouse_addresses[warehouse_index]
            if isinstance(warehouse_obj, dict):
                customer_warehouse_address = warehouse_obj.get("address", customer_warehouse_address)
            else:
                customer_warehouse_address = str(warehouse_obj)

    # 6b. Get signatory contact
    # Priority: signatory_contact_id > customer.general_director fields
    customer_signatory_name = customer.get("general_director_name") or ""
    customer_signatory_position = customer.get("general_director_position") or "Генеральный директор"

    if signatory_contact_id:
        # New approach: fetch from customer_contacts table
        contact_result = supabase_client.table("customer_contacts").select("*")\
            .eq("id", str(signatory_contact_id))\
            .eq("organization_id", str(org_id))\
            .execute()

        if contact_result.data and len(contact_result.data) > 0:
            contact = contact_result.data[0]
            # Build signatory name in short format: "Фамилия И. О."
            first_name = contact.get("name") or ""
            last_name = contact.get("last_name") or ""
            patronymic = contact.get("patronymic") or ""

            # Use short format for official documents
            customer_signatory_name = format_signatory_name_short(last_name, first_name, patronymic)

            # Use position field (make mandatory when is_signatory, fallback to default)
            customer_signatory_position = contact.get("position") or "Генеральный директор"

    # 7. Process products and calculate totals
    products = []
    total_quantity = 0
    total_amount_vat = Decimal("0.00")
    total_vat_from_sales = Decimal("0.00")  # Accumulate actual VAT from calculations

    for item in items_result.data:
        item_id = item.get("id")
        quantity = Decimal(str(item.get("quantity", 1)))

        # Get SALES prices from calculation results (not purchase price!)
        # Field names: sales_price_per_unit_with_vat, sales_price_total_with_vat
        phase_results = calc_results_by_item.get(item_id, {})
        unit_price_vat = Decimal(str(phase_results.get("sales_price_per_unit_with_vat", 0)))
        total_price_vat = Decimal(str(phase_results.get("sales_price_total_with_vat", 0)))
        item_vat = Decimal(str(phase_results.get("vat_from_sales", 0)))

        # Fallback: if no calculation results, use base_price_vat (purchase price)
        if unit_price_vat == 0 and total_price_vat == 0:
            unit_price_vat = Decimal(str(item.get("base_price_vat", 0)))
            total_price_vat = unit_price_vat * quantity
            # Calculate VAT as 22% of base (no calculation results available)
            item_vat = total_price_vat / Decimal("1.22") * Decimal("0.22")

        product_code = item.get("product_code") or ""
        brand = item.get("brand") or ""
        idn_sku = item.get("idn_sku") or ""

        # Round prices to 2 decimal places
        unit_price_vat = unit_price_vat.quantize(Decimal("0.01"))
        total_price_vat = total_price_vat.quantize(Decimal("0.01"))

        products.append({
            "name": item.get("product_name") or "",
            "product_code": product_code,
            "brand": brand,
            "idn_sku": idn_sku,
            "quantity": quantity,
            "unit_price_vat": unit_price_vat,
            "total_price_vat": total_price_vat
        })

        total_quantity += int(quantity)
        total_amount_vat += total_price_vat
        total_vat_from_sales += item_vat

    # 8. Use actual VAT from calculation results (not hardcoded 20%)
    vat_amount = total_vat_from_sales

    # Round totals to 2 decimal places
    total_amount_vat = total_amount_vat.quantize(Decimal("0.01"))
    vat_amount = vat_amount.quantize(Decimal("0.01"))

    # 9. Convert amounts to Russian words
    currency = quote.get("currency", "USD")
    total_amount_words = number_to_russian_words(total_amount_vat, currency)
    vat_amount_words = number_to_russian_words(vat_amount, currency)

    # 10. Build specification data
    specification_data = {
        # Contract info
        "contract_number": contract.get("contract_number", ""),
        "contract_date": contract.get("contract_date", ""),
        "specification_number": contract.get("next_specification_number", 1),
        "specification_date": date.today().isoformat(),

        # Company info
        # Use seller_company from seller_companies table if available, fallback to organization
        "seller_company": seller_company_data.get("name") or seller_company_name or organization.get("name", ""),
        "customer_name": customer.get("name", ""),

        # Quote info
        "quote_number": quote.get("idn_quote", ""),
        "currency": currency,
        # Use format_payment_terms from export_data_mapper (formats advance_from_client)
        "payment_terms": format_payment_terms(calculation_variables) if calculation_variables else quote.get("payment_terms", ""),
        # Use delivery_time from calculation variables (D9/B6 in Excel)
        "delivery_days": calculation_variables.get("delivery_time", quote.get("delivery_days", 30)),
        "delivery_terms": quote.get("delivery_terms", ""),

        # Products
        "products": products,

        # Totals
        "total_quantity": total_quantity,
        "total_amount_vat": total_amount_vat,
        "total_amount_words": total_amount_words,
        "vat_amount_words": vat_amount_words,

        # Addresses
        "customer_registration_address": customer.get("address", ""),
        "customer_warehouse_address": customer_warehouse_address,

        # Signatories
        # Use seller_company director if available, fallback to organization director
        "seller_signatory_position": (
            seller_company_data.get("general_director_position")
            or organization.get("general_director_position")
            or "Генеральный директор"
        ),
        "seller_signatory_name": (
            _format_seller_company_director_name(seller_company_data)
            or _format_org_director_name(organization)
        ),
        "customer_signatory_position": customer_signatory_position,
        "customer_signatory_name": customer_signatory_name,

        # Optional
        "additional_conditions": ""  # Can be provided by user
    }

    return specification_data


# ============================================================================
# DOCX GENERATION
# ============================================================================

def replace_variables_in_docx(doc: Document, variables: Dict[str, Any]) -> Document:
    """
    Replace template variables in DOCX document

    Variables format: [Variable Name]

    Args:
        doc: python-docx Document object
        variables: Dictionary of variable name -> value mappings

    Returns:
        Modified Document object
    """
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in variables.items():
            placeholder = f"[{key}]"
            if placeholder in paragraph.text:
                # Replace placeholder with value
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in variables.items():
                        placeholder = f"[{key}]"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))

    return doc


def remove_column(table, col_index: int):
    """
    Remove a column from a Word table by removing all cells at the given index.

    Note: This works by removing each cell individually, which effectively
    shifts all subsequent columns left.
    """
    for row in table.rows:
        # Get the cell to remove
        cell = row.cells[col_index]
        # Remove the cell's XML element
        tc = cell._tc
        tc.getparent().remove(tc)


def fill_products_table(doc: Document, products: List[Dict[str, Any]]) -> Document:
    """
    Fill products table in specification template

    Template has 8 columns:
    0. № (number)
    1. IDN-SKU - hidden if all blank
    2. Наименование позиции (name)
    3. Артикул (product_code) - hidden if all blank
    4. Производитель (brand) - hidden if all blank
    5. Количество (quantity)
    6. Цена + НДС (unit_price_vat) - SALES price, not purchase
    7. Стоимость + НДС (total_price_vat) - SALES total, not purchase

    Args:
        doc: python-docx Document object
        products: List of product dictionaries with sales prices

    Returns:
        Modified Document object
    """
    # Find products table (assume it's the first table with 8+ columns)
    products_table = None
    for table in doc.tables:
        if len(table.columns) >= 8:
            products_table = table
            break

    if not products_table:
        # If no table found, just return doc as-is
        return doc

    # Check if all idn_sku values are blank
    all_idn_sku_blank = all(not (product.get("idn_sku") or "").strip() for product in products)
    # Check if all product_code values are blank
    all_product_code_blank = all(not (product.get("product_code") or "").strip() for product in products)
    # Check if all brand values are blank
    all_brand_blank = all(not (product.get("brand") or "").strip() for product in products)

    # Add rows for each product (starting after header row)
    for idx, product in enumerate(products, start=1):
        row = products_table.add_row()

        # Fill cells - template column indices (before removal)
        row.cells[0].text = str(idx)  # №
        row.cells[1].text = product.get("idn_sku") or ""  # IDN-SKU
        row.cells[2].text = product.get("name") or ""
        row.cells[3].text = product.get("product_code") or ""
        row.cells[4].text = product.get("brand") or ""
        row.cells[5].text = str(int(product.get("quantity") or 0))
        row.cells[6].text = f"{product.get('unit_price_vat', 0):.2f}"
        row.cells[7].text = f"{product.get('total_price_vat', 0):.2f}"

    # Remove columns in reverse order (to preserve indices)
    # Track which columns to remove
    columns_to_remove = []

    # Remove IDN-SKU column if all blank (index 1)
    if all_idn_sku_blank:
        columns_to_remove.append(1)

    # Remove brand column if all blank (index 4)
    if all_brand_blank:
        columns_to_remove.append(4)

    # Remove product_code column if all blank (index 3)
    if all_product_code_blank:
        columns_to_remove.append(3)

    # Sort in reverse order to remove from right to left
    columns_to_remove.sort(reverse=True)

    # Remove columns
    for col_idx in columns_to_remove:
        try:
            remove_column(products_table, col_idx)
        except Exception:
            # If removal fails, continue with other columns
            pass

    return doc


async def generate_specification(
    quote_id: UUID,
    contract_id: UUID,
    org_id: UUID,
    user_id: UUID,
    additional_conditions: Optional[str],
    supabase_client,
    warehouse_index: Optional[int] = None,
    delivery_address_id: Optional[UUID] = None,
    signatory_contact_id: Optional[UUID] = None
) -> str:
    """
    Generate specification DOCX document

    Steps:
    1. Gather all data
    2. Load template DOCX
    3. Replace variables
    4. Fill products table
    5. Save to temporary file
    6. Return file path

    Args:
        quote_id: Quote UUID
        contract_id: Contract UUID
        org_id: Organization UUID
        user_id: User UUID (for audit trail)
        additional_conditions: Optional additional conditions text
        supabase_client: Supabase client
        warehouse_index: Legacy - Index into customer.warehouse_addresses JSONB array
        delivery_address_id: New - UUID of delivery address from customer_delivery_addresses table
        signatory_contact_id: New - UUID of signatory contact from customer_contacts table

    Returns:
        Path to generated DOCX file

    Raises:
        ValueError: If data is invalid or missing
        FileNotFoundError: If template file not found
    """

    # 1. Gather data
    data = await gather_specification_data(
        quote_id=quote_id,
        contract_id=contract_id,
        org_id=org_id,
        supabase_client=supabase_client,
        warehouse_index=warehouse_index,
        delivery_address_id=delivery_address_id,
        signatory_contact_id=signatory_contact_id
    )

    # Add additional conditions if provided
    if additional_conditions:
        data["additional_conditions"] = additional_conditions

    # 2. Load template
    template_path = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        "templates",
        "specification_template.docx"
    )

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")

    doc = Document(template_path)

    # 3. Prepare variables for replacement (exclude products list)
    variables = {k: v for k, v in data.items() if k != "products"}

    # 4. Replace variables
    doc = replace_variables_in_docx(doc, variables)

    # 5. Fill products table
    doc = fill_products_table(doc, data["products"])

    # 6. Save to temporary file
    output_dir = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        "temp_exports"
    )
    os.makedirs(output_dir, exist_ok=True)

    # Filename format: spec-{quote_idn}-{spec_number}.docx
    # Example: spec-КП25-0018-1.docx (first specification for quote КП25-0018)
    quote_idn = data.get("quote_number") or f"quote-{str(quote_id)[:8]}"
    spec_number = data["specification_number"]
    output_filename = f"spec-{quote_idn}-{spec_number}.docx"
    output_path = os.path.join(output_dir, output_filename)

    doc.save(output_path)

    # 7. Create audit record in specification_exports table
    # Convert Decimal values to strings for JSON serialization
    def convert_decimals(obj):
        if isinstance(obj, Decimal):
            return str(obj)
        elif isinstance(obj, dict):
            return {k: convert_decimals(v) for k, v in obj.items()}
        elif isinstance(obj, list):
            return [convert_decimals(item) for item in obj]
        return obj

    export_data_serializable = convert_decimals(data)

    export_record = {
        "organization_id": str(org_id),
        "quote_id": str(quote_id),
        "contract_id": str(contract_id),
        "specification_number": data["specification_number"],
        "specification_date": data["specification_date"],
        "export_data": export_data_serializable,  # Store snapshot of data
        "created_by": str(user_id)
    }

    supabase_client.table("specification_exports").insert(export_record).execute()

    # 8. Increment specification number in contract
    supabase_client.table("customer_contracts")\
        .update({"next_specification_number": data["specification_number"] + 1})\
        .eq("id", str(contract_id))\
        .execute()

    return output_path
