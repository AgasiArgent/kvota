"""
Create Basic Specification Template DOCX
Run this script once to generate the initial template file
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os


def create_specification_template():
    """Create a basic specification template with placeholders"""

    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("СПЕЦИФИКАЦИЯ")
    run.bold = True
    run.font.size = Pt(16)

    # Subtitle with contract info
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run(f"№ [specification_number] от [specification_date]\n")
    subtitle.add_run(f"к Договору № [contract_number] от [contract_date]")

    doc.add_paragraph()  # Spacing

    # Parties
    doc.add_paragraph(f"Продавец: [seller_company]")
    doc.add_paragraph(f"Покупатель: [customer_name]")
    doc.add_paragraph()

    # Main text
    doc.add_paragraph(
        "Настоящая Спецификация является неотъемлемой частью Договора поставки "
        "№ [contract_number] от [contract_date] г. и определяет перечень, количество, "
        "цены и стоимость товара, поставляемого Продавцом Покупателю."
    )
    doc.add_paragraph()

    # Products table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    headers = ["№", "IDN-SKU", "Наименование", "Артикул", "Производитель", "Кол-во", "Цена с НДС", "Стоимость с НДС"]
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True

    # Note: Products will be added programmatically

    doc.add_paragraph()

    # Totals
    doc.add_paragraph(f"Всего наименований: [total_quantity] шт.")
    doc.add_paragraph(f"Общая стоимость с НДС: [total_amount_vat] [currency]")
    doc.add_paragraph(f"Сумма прописью: [total_amount_words]")
    doc.add_paragraph(f"В том числе НДС: [vat_amount_words]")
    doc.add_paragraph()

    # Terms
    doc.add_paragraph("УСЛОВИЯ ПОСТАВКИ:")
    doc.add_paragraph(f"Условия оплаты: [payment_terms]")
    doc.add_paragraph(f"Срок поставки: [delivery_days] дней")
    doc.add_paragraph(f"Условия поставки: [delivery_terms]")
    doc.add_paragraph(f"Юридический адрес Покупателя: [customer_registration_address]")
    doc.add_paragraph(f"Адрес склада Покупателя: [customer_warehouse_address]")
    doc.add_paragraph()

    # Additional conditions
    doc.add_paragraph("Иные условия:")
    doc.add_paragraph("[additional_conditions]")
    doc.add_paragraph()
    doc.add_paragraph()

    # Signatures
    signatures_table = doc.add_table(rows=3, cols=2)
    signatures_table.style = 'Table Grid'

    # Row 1: Headers
    signatures_table.rows[0].cells[0].text = "ПРОДАВЕЦ"
    signatures_table.rows[0].cells[1].text = "ПОКУПАТЕЛЬ"

    # Row 2: Positions and names
    signatures_table.rows[1].cells[0].text = "[seller_signatory_position]\n[seller_signatory_name]"
    signatures_table.rows[1].cells[1].text = "[customer_signatory_position]\n[customer_signatory_name]"

    # Row 3: Signature lines
    signatures_table.rows[2].cells[0].text = "\n_______________ / _______________"
    signatures_table.rows[2].cells[1].text = "\n_______________ / _______________"

    # Save template
    template_path = os.path.join(
        os.path.dirname(__file__),
        "templates",
        "specification_template.docx"
    )

    doc.save(template_path)
    print(f"✓ Template created: {template_path}")


if __name__ == "__main__":
    create_specification_template()
